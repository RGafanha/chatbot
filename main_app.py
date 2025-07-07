import os
import io
import json
import logging
import base64
from datetime import datetime
from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv
import docx
from docx.shared import Inches, Pt, RGBColor

# --- Importações do LangChain ---
from langchain_community.document_loaders import PyPDFLoader, UnstructuredWordDocumentLoader, TextLoader
from langchain_community.vectorstores.utils import filter_complex_metadata
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough, RunnableParallel

# --- Configuração Inicial ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
load_dotenv()

from fastapi.staticfiles import StaticFiles

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")

# --- Variáveis Globais ---
retriever = None
qa_chain = None

# =============================================================================
# SEÇÃO 1: LÓGICA DE NEGÓCIO (Restaurada e Adaptada para FastAPI)
# =============================================================================

def load_and_process_documents():
    # (Mesma lógica de antes)
    if not os.path.exists("documentos"):
        os.makedirs("documentos")
        return None
    PERSIST_DIRECTORY = "chroma_db_chatbot_qa"
    if not os.path.exists(PERSIST_DIRECTORY):
        all_loaded_documents = []
        for filename in os.listdir("documentos"):
            file_path = os.path.join("documentos", filename)
            loader = None
            if filename.lower().endswith(".pdf"): loader = PyPDFLoader(file_path)
            elif filename.lower().endswith((".docx", ".doc")): loader = UnstructuredWordDocumentLoader(file_path)
            elif filename.lower().endswith(".txt"): loader = TextLoader(file_path, encoding='utf-8')
            if loader:
                try: all_loaded_documents.extend(loader.load())
                except Exception: pass
        if not all_loaded_documents: return None
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_documents(all_loaded_documents)
        chunks_com_metadata_simples = filter_complex_metadata(chunks)
        vectordb = Chroma.from_documents(documents=chunks_com_metadata_simples, embedding=OpenAIEmbeddings(), persist_directory=PERSIST_DIRECTORY)
    else:
        vectordb = Chroma(persist_directory=PERSIST_DIRECTORY, embedding_function=OpenAIEmbeddings())
    return vectordb.as_retriever(search_kwargs={"k": 3})

def get_qa_chain(_retriever):
    # (Mesma lógica de antes)
    llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0)
    prompt_template = "Você é um assistente de IA focado em responder perguntas com base em um contexto fornecido. Sua tarefa é usar apenas as informações dos documentos abaixo para formular sua resposta. Se a resposta não puder ser encontrada nos documentos, você deve informar explicitamente: 'Infelizmente, não foi encontrado as informações solicitada.'. Não adicione nenhuma informação que não esteja nos trechos fornecidos.\n\nCONTEXTO: {context}\nPERGUNTA: {question}\nRESPOSTA:"
    PROMPT = PromptTemplate.from_template(prompt_template)
    def format_docs(docs):
        return "\n\n".join(f"--- Trecho de {os.path.basename(doc.metadata.get('source', ''))} ---\n{doc.page_content}" for doc in docs)
    return RunnableParallel({"context": _retriever, "question": RunnablePassthrough()}).assign(answer=(RunnablePassthrough.assign(context=(lambda x: format_docs(x["context"]))) | PROMPT | llm | StrOutputParser()))

def adicionar_documento_ao_banco(caminho_arquivo):
    # (Lógica de adicionar ao ChromaDB)
    PERSIST_DIRECTORY = "chroma_db_chatbot_qa"
    if caminho_arquivo.lower().endswith(".pdf"): loader = PyPDFLoader(caminho_arquivo)
    elif caminho_arquivo.lower().endswith((".docx", ".doc")): loader = UnstructuredWordDocumentLoader(caminho_arquivo)
    elif caminho_arquivo.lower().endswith(".txt"): loader = TextLoader(caminho_arquivo, encoding='utf-8')
    else: return
    documentos = loader.load()
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(documentos)
    chunks = filter_complex_metadata(chunks)
    vectordb = Chroma(persist_directory=PERSIST_DIRECTORY, embedding_function=OpenAIEmbeddings())
    vectordb.add_documents(chunks)
    vectordb.persist()

def enhance_procedure_with_ai(procedure_text):
    # (Lógica de aprimoramento)
    llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.3)
    try:
        with open("aprimoramento_template.txt", 'r', encoding='utf-8') as f:
            enhancement_prompt_template = f.read()
    except FileNotFoundError:
        return "Erro: O arquivo de template 'aprimoramento_template.txt' não foi encontrado."
    prompt = PromptTemplate.from_template(enhancement_prompt_template)
    chain = prompt | llm | StrOutputParser()
    return chain.invoke({"texto_do_procedimento": procedure_text})

def extract_text_and_images_from_bytes(file_bytes):
    # (Lógica de extração de texto)
    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join([para.text for para in doc.paragraphs])

def reformat_document_to_json(content_antigo, regras_padrao, data_atual):
    # (Lógica de reformatar para JSON)
    llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.0, response_format={"type": "json_object"})
    prompt_template = """Siga as regras para analisar o conteúdo antigo e retornar um objeto JSON válido. A versão do novo documento será sempre '1.0'. A data do novo documento será a data atual fornecida. Data Atual: {data}. REGRAS: {regras}. CONTEÚDO DO DOCUMENTO ANTIGO: {conteudo}"""
    prompt = PromptTemplate.from_template(prompt_template)
    chain = prompt | llm | StrOutputParser()
    response_str = chain.invoke({"data": data_atual, "regras": regras_padrao, "conteudo": content_antigo})
    return json.loads(response_str)

def create_documentation_from_scratch(author_name, raw_text, data_atual, creation_prompt_template):
    # (Lógica de criar do zero)
    llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.2, response_format={"type": "json_object"})
    prompt = PromptTemplate.from_template(creation_prompt_template)
    chain = prompt | llm | StrOutputParser()
    response_str = chain.invoke({"nome_do_autor": author_name, "data_atual": data_atual, "texto_bruto_usuario": raw_text})
    return json.loads(response_str)

def create_docx_in_memory(structured_data, images_data=None):
    logging.info("--- Iniciando create_docx_in_memory ---")
    new_doc = docx.Document()
    header = structured_data.get("header_data", {})
    body = structured_data.get("body_content", {})
    # ... (código da tabela do cabeçalho permanece o mesmo)
    table = new_doc.add_table(rows=3, cols=4); table.style = 'Table Grid'
    cell_logo = table.cell(0, 0); p_logo = cell_logo.paragraphs[0]; run_logo = p_logo.add_run()
    try: run_logo.add_picture('logo.png', width=Inches(1.0))
    except FileNotFoundError: p_logo.text = "Wilson.Sons"
    cell_proc = table.cell(0, 1); cell_proc.text = header.get("procedimento_descricao", ""); cell_proc.merge(table.cell(0, 2))
    p_codigo = table.cell(0, 3).paragraphs[0]; p_codigo.add_run('Código: ').bold = True; p_codigo.add_run(header.get('codigo', ''))
    cell_title = table.cell(1, 0); p_title = cell_title.paragraphs[0]; p_title.add_run('Título: ').bold = True; p_title.add_run(header.get('titulo', '')); cell_title.merge(table.cell(1, 1)).merge(table.cell(1, 2)).merge(table.cell(1, 3))
    p_versao = table.cell(2, 0).paragraphs[0]; p_versao.add_run('Versão: ').bold = True; p_versao.add_run(header.get('versao', ''))
    p_emitentes = table.cell(2, 1).paragraphs[0]; p_emitentes.add_run('Emitentes: ').bold = True; p_emitentes.add_run(header.get('emitentes', ''))
    p_aprovador = table.cell(2, 2).paragraphs[0]; p_aprovador.add_run('Aprovador: ').bold = True; p_aprovador.add_run(header.get('aprovador', ''))
    p_data = table.cell(2, 3).paragraphs[0]; p_data.add_run('Data: ').bold = True; p_data.add_run(header.get('data', ''))
    new_doc.add_paragraph()

    sorted_body_keys = sorted(body.keys())
    for section_key in sorted_body_keys:
        section_content = body.get(section_key, "")
        logging.info(f"Processando seção do corpo: '{section_key}'")
        try:
            num, title_text = section_key.split('-', 1)
            title_text = title_text.replace('_', ' ')
            heading = new_doc.add_heading(f"{num}. {title_text}", level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        except ValueError:
            new_doc.add_heading(section_key, level=2).runs[0].font.color.rgb = RGBColor(0, 0, 0)

        content_list = section_content if isinstance(section_content, list) else [str(section_content)]
        import re
        for item in content_list:
            logging.info(f"Processando item de conteúdo: {item}") # Log de DEBUG para INFO
            parts = re.split(r'(\[\s*IMAGEM[_-]?(\d+)\s*\])', item, flags=re.IGNORECASE)
            for part in parts:
                if not part: continue
                match = re.match(r'\[\s*IMAGEM[_-]?(\d+)\s*\]', part, flags=re.IGNORECASE)
                if match and images_data:
                    img_index = match.group(1)
                    img_key = f"image_{img_index}"
                    logging.info(f"Placeholder '{part.strip()}' encontrado. Procurando por chave de imagem '{img_key}'.")
                    if img_key in images_data:
                        logging.info(f"Chave '{img_key}' encontrada. Inserindo imagem.")
                        img_bytes = images_data[img_key]
                        try:
                            new_doc.add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
                            new_doc.add_paragraph()
                        except Exception as e:
                            logging.error(f"Falha ao inserir imagem {img_key} no DOCX: {e}")
                            new_doc.add_paragraph(f"[Erro ao carregar imagem {img_index}]", style='Comment')
                    else:
                        logging.warning(f"Chave de imagem '{img_key}' NÃO encontrada no dicionário de imagens. Placeholder será inserido como texto.")
                        new_doc.add_paragraph(part)
                elif part.strip():
                    logging.info(f"Adicionando parte de texto: {part[:100]}...") # Log de DEBUG para INFO
                    new_doc.add_paragraph(part)

    # Aplica a fonte Arial em todo o documento antes de salvar
    for paragraph in new_doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'

    for table in new_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'

    file_stream = io.BytesIO()
    new_doc.save(file_stream)
    file_stream.seek(0)
    logging.info("--- Finalizando create_docx_in_memory ---")
    return file_stream.getvalue()

def get_internet_answer_chain():
    llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.5)
    prompt_template = """
    Você é um assistente de pesquisa IA. Sua tarefa é gerar uma resposta detalhada e bem estruturada para a pergunta do usuário, simulando um resultado de uma pesquisa aprofundada na internet.
    Se a pergunta for sobre como realizar uma tarefa técnica (como uma instalação ou configuração), forneça uma resposta em formato de passo a passo.
    Seja claro, completo e profissional.

    PERGUNTA DO USUÁRIO:
    {question}

    RESPOSTA DETALHADA (simulando pesquisa na web):
    """
    prompt = PromptTemplate.from_template(prompt_template)
    chain = prompt | llm | StrOutputParser()
    return chain


# =============================================================================
# SEÇÃO 2: API ENDPOINTS (Atualizados e Expandidos)
# =============================================================================

@app.on_event("startup")
async def startup_event():
    # (Mesma lógica de antes)
    global retriever, qa_chain
    if not os.getenv("OPENAI_API_KEY"): logging.error("Chave da API da OpenAI não encontrada!"); return
    try:
        retriever = load_and_process_documents()
        if retriever: qa_chain = get_qa_chain(retriever); logging.info("Chatbot pronto.")
        else: logging.warning("Base de conhecimento vazia.")
    except Exception as e: logging.error(f"Erro na inicialização: {e}")

@app.get("/", response_class=HTMLResponse)
async def read_root():
    with open("index.html", "r", encoding="utf-8") as f: return HTMLResponse(content=f.read())

class AskRequest(BaseModel):
    question: str

class WebAskRequest(BaseModel):
    question: str
    author: str

@app.post("/api/ask")
async def api_ask(request: AskRequest):
    # (Mesma lógica de antes)
    if not qa_chain: return JSONResponse(status_code=500, content={"error": "Chatbot não inicializado."})
    try:
        result = qa_chain.invoke(request.question)
        sources = [os.path.basename(doc.metadata.get('source', 'N/A')) for doc in result["context"]]
        return {"answer": result["answer"], "sources": sorted(list(set(sources)))}
    except Exception as e: return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/api/ask_web")
async def api_ask_web(request: WebAskRequest):
    """Recebe uma pergunta, busca na web (simulado) e retorna o conteúdo para criar um documento."""
    try:
        # 1. Gerar o conteúdo usando o LLM para simular a pesquisa
        logging.info(f"Iniciando pesquisa na web para a pergunta: {request.question}")
        web_search_chain = get_internet_answer_chain()
        web_content = web_search_chain.invoke({"question": request.question})
        logging.info(f"Conteúdo gerado pela pesquisa na web: {web_content[:200]}...")

        # 2. Usar o conteúdo gerado para criar a estrutura do documento
        with open("criacao_template.txt", 'r', encoding='utf-8') as f:
            regras_criacao = f.read()
        
        now = datetime.now()
        date_for_document = now.strftime("%d/%m/%Y")

        structured_data = create_documentation_from_scratch(
            author_name=request.author,
            raw_text=web_content,
            data_atual=date_for_document,
            creation_prompt_template=regras_criacao
        )
        logging.info(f"Dados estruturados criados a partir do conteúdo da web: {json.dumps(structured_data, indent=2)}")

        # 3. Gerar o arquivo .docx em memória
        final_doc_bytes = create_docx_in_memory(structured_data, images_data=None)

        # 4. Preparar o nome do arquivo e o conteúdo para a resposta
        doc_title = structured_data.get("header_data", {}).get("titulo", "Novo_Documento").replace(" ", "_")
        new_filename = f"PROC_TI_{now.strftime('%d_%m_%Y')}_{doc_title}.docx"
        
        return {
            "message": "Pesquisa na web concluída e documento preliminar gerado.",
            "web_content": web_content,
            "filename": new_filename,
            "file_content_b64": base64.b64encode(final_doc_bytes).decode('utf-8')
        }
    except Exception as e:
        logging.error(f"Erro na API /api/ask_web: {e}", exc_info=True)
        return JSONResponse(status_code=500, content={"error": str(e)})


class DocumentData(BaseModel):
    file_content_b64: str
    filename: str

@app.post("/api/add_to_kb")
async def api_add_to_kb(data: DocumentData):
    """Adiciona um documento .docx (recebido como base64) à base de conhecimento."""
    try:
        os.makedirs("documentos", exist_ok=True)
        file_bytes = base64.b64decode(data.file_content_b64)
        output_path = os.path.join("documentos", data.filename)
        with open(output_path, "wb") as f: f.write(file_bytes)
        adicionar_documento_ao_banco(output_path)
        return {"message": f"Documento '{data.filename}' adicionado com sucesso à base de conhecimento!"}
    except Exception as e: return JSONResponse(status_code=500, content={"error": str(e)})

class EnhanceRequest(BaseModel):
    procedure_text: str

class ApplyEnhancementRequest(BaseModel):
    original_file_content_b64: str
    enhancement_suggestions: str
    original_filename: str

@app.post("/api/enhance")
async def api_enhance(request: EnhanceRequest):
    """Recebe um texto de procedimento e retorna sugestões de aprimoramento."""
    try:
        suggestions = enhance_procedure_with_ai(request.procedure_text)
        return {"suggestions": suggestions}
    except Exception as e: return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/api/apply_enhancement")
async def api_apply_enhancement(request: ApplyEnhancementRequest):
    """Aplica as sugestões de aprimoramento a um documento e retorna o novo arquivo."""
    try:
        logging.info("Iniciando a aplicação do aprimoramento...")
        
        # Decodifica o conteúdo do arquivo original
        original_file_bytes = base64.b64decode(request.original_file_content_b64)
        
        # Extrai o texto do documento original para obter o contexto completo
        original_text = extract_text_and_images_from_bytes(original_file_bytes)

        # Usa a IA para reescrever o procedimento com base nas sugestões
        llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.1, response_format={"type": "json_object"})
        rewrite_prompt_template = """Você é um Redator Técnico especialista. Sua tarefa é reescrever o 'PROCEDIMENTO ORIGINAL' de um documento, aplicando as 'SUGESTÕES DE APRIMORAMENTO' fornecidas. Mantenha o restante do documento intacto. Retorne um objeto JSON contendo o texto reescrito na chave 'procedimento_aprimorado'.\n\nPROCEDIMENTO ORIGINAL:\n{procedimento_original}\n\nSUGESTÕES DE APRIMORAMENTO:\n{sugestoes}\n\nJSON DE SAÍDA:\n{{"procedimento_aprimorado": "[Seu texto do procedimento reescrito aqui]"}}"""
        prompt = PromptTemplate.from_template(rewrite_prompt_template)
        chain = prompt | llm | StrOutputParser()
        
        response_str = chain.invoke({
            "procedimento_original": original_text,
            "sugestoes": request.enhancement_suggestions
        })
        
        rewritten_data = json.loads(response_str)
        enhanced_procedure = rewritten_data.get("procedimento_aprimorado", "")

        # Recria a estrutura do documento com o procedimento aprimorado
        with open("padrao_template.txt", 'r', encoding='utf-8') as f: regras_padrao = f.read()
        now = datetime.now()
        date_for_document = now.strftime("%d/%m/%Y")
        
        # Reestrutura o documento inteiro para garantir a consistência
        structured_data = reformat_document_to_json(enhanced_procedure, regras_padrao, date_for_document)

        # Gera o novo arquivo .docx
        final_doc_bytes = create_docx_in_memory(structured_data, original_file_bytes)
        
        # Mantém o nome do arquivo original
        new_filename = request.original_filename

        return {
            "message": "Documento aprimorado com sucesso!",
            "filename": new_filename,
            "file_content_b64": base64.b64encode(final_doc_bytes).decode('utf-8')
        }

    except Exception as e:
        logging.error(f"Erro ao aplicar aprimoramento: {e}", exc_info=True)
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/api/standardize")
async def api_standardize(file: UploadFile = File(...)):
    """Padroniza um DOCX e retorna o resultado e o texto do procedimento."""
    try:
        with open("padrao_template.txt", 'r', encoding='utf-8') as f: regras_padrao = f.read()
        now = datetime.now()
        date_for_document = now.strftime("%d/%m/%Y")
        original_filename = file.filename
        file_bytes = await file.read()
        
        conteudo_antigo_texto = extract_text_and_images_from_bytes(file_bytes)
        structured_data = reformat_document_to_json(conteudo_antigo_texto, regras_padrao, date_for_document)
        
        final_doc_bytes = create_docx_in_memory(structured_data, file_bytes)
        
        new_filename = f"PROC_TI_{now.strftime('%d_%m_%Y')}_{original_filename}"
        procedure_text = structured_data.get("body_content", {}).get("4-Procedimentos", "")

        return {
            "message": "Documento padronizado com sucesso!",
            "filename": new_filename,
            "file_content_b64": base64.b64encode(final_doc_bytes).decode('utf-8'),
            "procedure_text": procedure_text
        }
    except Exception as e: return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/api/create")
async def api_create(
    author: str = Form(...),
    description: str = Form(...),
    image_1: UploadFile = File(None),
    image_2: UploadFile = File(None),
    image_3: UploadFile = File(None)
):
    """Cria um novo documento a partir de texto e imagens."""
    logging.info("--- Iniciando /api/create ---")
    try:
        logging.info(f"Form data received for author: '{author}' and description: '{description}'")

        images_data = {}
        possible_images = {"image_1": image_1, "image_2": image_2, "image_3": image_3}
        for key, ufile in possible_images.items():
            if ufile and ufile.filename:
                logging.info(f"Found UploadFile with key: '{key}' and filename: '{ufile.filename}'")
                image_bytes = await ufile.read()
                if image_bytes:
                    images_data[key] = image_bytes
                else:
                    logging.warning(f"Image file '{key}' is empty.")
            
        logging.info(f"Final images_data dictionary has keys: {list(images_data.keys())}")

        with open("criacao_template.txt", 'r', encoding='utf-8') as f: 
            regras_criacao = f.read()
        
        now = datetime.now()
        date_for_document = now.strftime("%d/%m/%Y")

        structured_data = create_documentation_from_scratch(author, description, date_for_document, regras_criacao)
        logging.info(f"Dados estruturados recebidos da IA: {json.dumps(structured_data, indent=2)}")

        final_doc_bytes = create_docx_in_memory(structured_data, images_data=images_data)
        
        doc_title = structured_data.get("header_data", {}).get("titulo", "Novo_Documento").replace(" ", "_")
        new_filename = f"PROC_TI_{now.strftime('%d_%m_%Y')}_{doc_title}.docx"

        logging.info(f"--- Finalizando /api/create com sucesso ---")
        return {
            "message": "Documento criado com sucesso!",
            "filename": new_filename,
            "file_content_b64": base64.b64encode(final_doc_bytes).decode('utf-8')
        }
    except Exception as e:
        logging.error(f"Erro na criação do documento: {e}", exc_info=True)
        return JSONResponse(status_code=500, content={"error": str(e)})

# --- Execução do Servidor ---
if __name__ == "__main__":
    import uvicorn
    if not os.getenv("OPENAI_API_KEY"): print("AVISO: Chave da API da OpenAI não encontrada.")
    uvicorn.run(app, host="0.0.0.0", port=8000)